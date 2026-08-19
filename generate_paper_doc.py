import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_document():
    doc = docx.Document()

    # Set page margins to standard 0.8 inch for clean 2-page layout
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Styling helper functions
    def set_font(run, font_name="Calibri", size_pt=11, bold=False, italic=False, color_rgb=(0,0,0)):
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, font_name="Calibri", size_pt=14, bold=True, color_rgb=(24, 76, 120))
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        set_font(run, font_name="Calibri", size_pt=12, bold=True, color_rgb=(40, 40, 40))
        return p

    def add_body_p(text, bold_prefix=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            set_font(run_b, font_name="Calibri", size_pt=10.5, bold=True, color_rgb=(30, 30, 30))
        run_t = p.add_run(text)
        set_font(run_t, font_name="Calibri", size_pt=10.5, bold=False, color_rgb=(50, 50, 50))
        return p

    def add_bullet(p_or_text, bold_title="", text_body=""):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        if bold_title:
            r_b = p.add_run(bold_title)
            set_font(r_b, font_name="Calibri", size_pt=10.5, bold=True, color_rgb=(20, 20, 20))
        r_t = p.add_run(text_body)
        set_font(r_t, font_name="Calibri", size_pt=10.5, bold=False, color_rgb=(50, 50, 50))
        return p

    # --- TITLE & HEADER BLOCK ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_t = p_title.add_run("VERITAS AI: REAL-TIME AUTOMATED FACT-CHECKING SYSTEM")
    set_font(r_t, font_name="Calibri", size_pt=16, bold=True, color_rgb=(15, 45, 85))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run("Core Research Components for Academic Paper Publication: Motivation, Research Gap, Objectives, and Novelty")
    set_font(r_sub, font_name="Calibri", size_pt=11, italic=True, color_rgb=(100, 100, 100))

    # --- PAGE 1: MOTIVATION & RESEARCH GAP ---

    # 1. MOTIVATION
    add_heading_1("1. MOTIVATION")
    add_body_p(
        "The hyper-proliferation of digital news across social platforms, messaging networks, and online portals has accelerated the rapid spread of misinformation, fabricated news, and deep-fake claims. Unverified viral statements regarding political leaders, public health guidelines, economic policies, and geopolitical events cause severe societal disruption, market volatility, and public panic within minutes of release."
    )
    add_body_p(
        "While professional fact-checking organizations (e.g., PIB Fact Check, Reuters, AltNews, Snopes) manually investigate claims, manual verification requires hours or days—creating a dangerous temporal window during which fake news spreads unchecked. Furthermore, existing automated artificial intelligence solutions either rely on static pre-trained dataset classification (which cannot verify breaking real-time news) or suffer from severe search bias, single-query retrieval failure, and high computational API costs."
    )
    add_body_p(
        "Therefore, there is an urgent research imperative to build an intelligent, sub-second, dynamic automated fact-checking framework capable of real-time multi-query web retrieval, semantic persistence memory, claim-evidence stance verification, and multi-provider auto-healing resilience."
    )

    # 2. RESEARCH GAP
    add_heading_1("2. RESEARCH GAP")
    add_body_p("A rigorous review of existing literature and automated fact-checking systems reveals four major critical research gaps:")

    add_bullet("", "Gap 1: Static Dataset Limitations and Absence of Real-Time Information Access – ", 
               "Traditional machine learning and static Large Language Model (LLM) classifiers (e.g., fine-tuned BERT, RoBERTa) evaluate claims based solely on static pre-training weights. They cannot verify post-cutoff events, breaking news, or newly emerging rumors, resulting in severe hallucinations or outright incorrect verdicts on contemporary claims.")

    add_bullet("", "Gap 2: Single-Query Search Bias and Query Drift in Retrieval-Augmented Generation (RAG) – ", 
               "Existing RAG-based fact-checkers take the raw, biased user claim (e.g., 'The president of america is Vijay') and pass it directly to web search engines. This introduces confirmation bias, as search engines return topically related articles about the candidate entity rather than retrieving objective counter-evidence regarding the actual office-holder.")

    add_bullet("", "Gap 3: Redundant Computational Latency and API Cost Inefficiency – ", 
               "Standard dynamic RAG pipelines re-execute expensive multi-step web retrieval, snippet processing, and LLM reasoning for every incoming user request—even when identical or paraphrased claims have already been verified minutes earlier. They lack a persistent, dynamic semantic knowledge base with incremental vector indexing.")

    add_bullet("", "Gap 4: Inadequate Stance Classification and Entity/Temporal Contradiction Handling – ", 
               "Current automated verification models struggle to distinguish between topically relevant background text and explicit refutation. When evidence mentions a different entity holding an exclusive role (e.g., 'Narendra Modi is Prime Minister') or indicates a planned future date (e.g., 'Chandrayaan-4 scheduled for 2027'), traditional models often default to 'UNCERTAIN' rather than correctly identifying a 'FAKE/REFUTED' contradiction.")

    # --- PAGE BREAK FOR FRONT & BACK PRINTING ---
    doc.add_page_break()

    # --- PAGE 2: OBJECTIVES & NOVELTY ---

    # 3. OBJECTIVES
    add_heading_1("3. SYSTEM OBJECTIVES")
    add_body_p("To resolve the aforementioned research gaps, the primary objective of this project is to design, implement, and evaluate VERITAS AI—an ultra-fast, multi-query, stance-aware automated fact-checking system equipped with persistent semantic memory and auto-healing multi-provider fallback. The specific technical objectives are as follows:")

    add_bullet("", "Objective 1: Build Sub-Second Semantic Persistence Memory – ", 
               "Integrate a persistent FAISS vector index with deterministic local subword embeddings to store verified fact-check records, enabling instantaneous (0.00s / <20ms) verdict retrieval for exact, paraphrased, or typo-containing claims.")

    add_bullet("", "Objective 2: Develop Bias-Free Multi-Query Intent Expansion – ", 
               "Engineer an intelligent query optimization module that expands raw user inputs into complementary search strategies (direct keywords, official government records, and targeted office-holder queries) to eliminate retrieval bias.")

    add_bullet("", "Objective 3: Implement Category-Aware Claim-Evidence Stance Verification – ", 
               "Formulate a multi-stage verification engine that categorizes claims into functional taxonomies (CURRENT_ROLE, TEMPORAL, NUMERICAL, EVENT) and evaluates source stances into SUPPORTS, REFUTES, or INSUFFICIENT using exclusive role and temporal contradiction rules.")

    add_bullet("", "Objective 4: Ensure High-Availability Multi-Provider LLM Resilience – ", 
               "Construct an auto-healing LLM orchestration client with automatic multi-provider fallback (Groq Cloud → Google Gemini Flash → OpenRouter) to guarantee uninterrupted service during rate-limit breaches or API outages.")

    # 4. NOVELTY & CORE CONTRIBUTIONS
    add_heading_1("4. NOVELTY & CORE CONTRIBUTIONS")
    add_body_p("The Veritas AI framework introduces four groundbreaking novel contributions to the domain of automated fact-checking and natural language processing:")

    add_bullet("", "Contribution 1: Deterministic SHA256 Subword Local Embedding Engine – ", 
               "Unlike conventional systems relying on network-dependent embedding APIs, Veritas AI introduces a fast, local, 1536-dimensional SHA256 character 3-gram subword hashing embedder. Running in 0.1ms with zero API cost, it provides 100% process-independent persistence and extreme resilience against spelling typos, word-order reordering, and phrasing variations.")

    add_bullet("", "Contribution 2: Role-Targeted Query Expansion for Entity-Bias Elimination – ", 
               "Veritas AI introduces an automated regex and syntactic role-parser that detects exclusive office-holder claims (e.g., PM, CM, President, CEO) and dynamically generates targeted queries (e.g., 'who is current president of america official'). This guarantees the retrieval of authoritative counter-evidence naming the actual office-holder.")

    add_bullet("", "Contribution 3: Dual-Stage Stance Verification with Fail-Safe Entity Extraction – ", 
               "The system incorporates a hybrid stance evaluation pipeline that combines structured LLM reasoning with a deterministic rule-based fallback safeguard. Even in the event of total network failure, the rule-based fallback inspects candidate snippets, detects official role holders, and synthesizes informative refutations naming the exact real-world figure.")

    add_bullet("", "Contribution 4: Entity-Compatible Incremental Semantic Knowledge Base – ", 
               "Veritas AI features an upgraded FAISS (v2.1) vector store backed by numeric entity compatibility guardrails. It prevents false matches between distinct quantitative figures (e.g., ₹2000 vs ₹500 notes) while enabling instant 0.00s verification reuse for semantic paraphrases across user sessions.")

    # Save document
    target_path = r"c:\My_Learning\FND Modified\Veritas_AI_Research_Paper_Components.docx"
    doc.save(target_path)
    print(f"SUCCESS: Research paper document generated successfully at: {target_path}")

if __name__ == "__main__":
    create_document()
